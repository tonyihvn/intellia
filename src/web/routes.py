from flask import Blueprint, request, jsonify, current_app, send_file
from ..core.query_handler import QueryHandler
from ..core.command_handler import CommandHandler
from ..llm.client import LLMClient
from ..rag.enhancer import QueryEnhancer
from ..rag.manager import RAGManager
from ..rag.schema_context import SchemaContextManager
from ..config import Config
from .history_manager import HistoryManager
import os
import logging
# Prefer mysql.connector if available, fall back to pymysql, otherwise leave a None placeholder
# so the code can report a helpful error when DB connector is not installed.
try:
    import mysql.connector as mysql_connector
except Exception:
    try:
        import pymysql as mysql_connector
    except Exception:
        mysql_connector = None
        logging.warning("Neither 'mysql.connector' nor 'pymysql' is installed; DB connector functionality will be unavailable. Install 'mysql-connector-python' or 'PyMySQL'.")

import re
import json
import html
from datetime import datetime
import uuid
from ..db.connection import get_db_connection
from ..rag.manager import RAGManager
from ..db.schema_fetcher import SchemaFetcher
import shutil
from pathlib import Path
import re
import base64
from io import BytesIO
try:
    from docx import Document
except Exception:
    Document = None

main_routes = Blueprint('main', __name__, url_prefix='/api')

def get_query_handler():
    """
    Create and return a properly initialized QueryHandler instance
    """
    try:
        # Get LLM configuration and return a ready QueryHandler
        llm_config = Config.get_llm_config()
        rag_manager = RAGManager()

        # Try to set DB context for RAG so schema snippets are available to QueryEnhancer
        try:
            conn = get_db_connection()
            if conn:
                try:
                    rag_manager.set_db_context(conn)
                    # Bootstrap RAG knowledge if empty
                    if rag_manager.is_empty():
                        try:
                            rag_manager.bootstrap_from_db(conn)
                        except Exception:
                            logging.warning('Failed to bootstrap RAG from DB during handler init')
                finally:
                    try:
                        conn.close()
                    except Exception:
                        pass
        except Exception:
            logging.warning('Could not initialize DB context for RAGManager in request')

        query_enhancer = QueryEnhancer(rag_manager)
        query_handler = QueryHandler(LLMClient(), query_enhancer)
        logging.info("Successfully initialized QueryHandler")
        return query_handler
    except Exception as e:
        logging.error(f"Error initializing QueryHandler: {e}")
        raise


def validate_sql_against_rag_schema(sql: str, db_conn):
    """Validate that SQL references only tables/columns present in the current DB schema.

    Returns (True, {}) on success or (False, details) where details contains unknown identifiers.
    Handles table aliases, subqueries, and common SQL patterns while validating table and column references.
    """
    try:
        # Common SQL keywords and functions that shouldn't trigger validation errors
        sql_keywords = {
            'select', 'from', 'where', 'and', 'or', 'join', 'inner', 'outer', 'left', 'right',
            'group', 'by', 'having', 'order', 'limit', 'offset', 'as', 'on', 'case', 'when',
            'then', 'else', 'end', 'union', 'all', 'in', 'exists', 'between', 'like', 'is',
            'null', 'not', 'true', 'false', 'asc', 'desc', 'distinct', 'count', 'sum', 'avg',
            'min', 'max', 'coalesce', 'if', 'ifnull', 'cast'
        }

        fetcher = SchemaFetcher(db_conn)
        cols = fetcher.get_columns()  # list of dicts with table_name and column_name
        table_columns = {}
        all_columns = set()
        all_tables = set()
        for c in cols:
            t = c.get('table_name')
            col = c.get('column_name') or c.get('column') or c.get('Field')
            if not t or not col:
                continue
            all_tables.add(t)
            table_columns.setdefault(t, set()).add(col)
            all_columns.add(col)

        # Add common English stopwords and collect table aliases to ignore in table validation
        stopwords = {'the', 'a', 'an', 'of', 'to', 'for', 'in', 'on', 'by', 'with', 'as', 'at', 'from'}
        # Find table aliases in FROM/JOIN clauses
        alias_pattern = r'(?:from|join)\s+([`\"]?\w+[`\"]?)(?:\s+(?:as\s+)?([`\"]?\w+[`\"]?))?'
        aliases = set()
        for match in re.finditer(alias_pattern, sql, re.IGNORECASE):
            table_name = match.group(1)
            alias = match.group(2)
            if alias:
                aliases.add(alias.strip('`"').lower())

        unknown = {'table_column': [], 'unqualified_columns': []}

        # Check occurrences of table.column (with optional quotes/backticks)
        for m in re.finditer(r"([`\"]?)([A-Za-z0-9_]+)\1\s*\.\s*([`\"]?)([A-Za-z0-9_]+)\3", sql):
            table = m.group(2)
            column = m.group(4)
            table_lc = table.lower()
            if (table_lc not in {t.lower() for t in all_tables}
                and table_lc not in stopwords
                and table_lc not in aliases) or \
               (column != '*' and column.lower() not in {c.lower() for c in table_columns.get(table, set())}):
                if table_lc not in stopwords and table_lc not in aliases:
                    unknown['table_column'].append({'table': table, 'column': column})

        # More lenient unqualified column check for SELECT clause
        sel = re.search(r"select\s+(.*?)\s+from\s", sql, flags=re.I | re.S)
        if sel:
            select_text = sel.group(1)
            # Remove function calls, subqueries and literals
            cleaned = re.sub(r"\b\w+\s*\([^)]*\)", '', select_text)  # Remove function calls
            cleaned = re.sub(r'\([^)]+\)', '', cleaned)  # Remove subqueries
            cleaned = re.sub(r"'[^']*'", '', cleaned)  # Remove string literals
            cleaned = re.sub(r'"[^"]*"', '', cleaned)  # Remove double-quoted strings
            # split by commas and extract potential column tokens
            parts = [p.strip() for p in cleaned.split(',') if p.strip()]
            
            for p in parts:
                # ignore wildcards, literals and known SQL keywords
                # Skip various cases that don't need validation
                if any([
                    p == '*' or p.endswith('.*'),  # wildcards
                    re.match(r"^\d+$", p),  # numbers
                    '.' in p,  # already handled qualified columns
                    not p,  # empty strings
                    p.lower() in sql_keywords,  # SQL keywords
                    re.match(r"^['\"].*['\"]$", p),  # string literals
                    re.match(r"^`.*`$", p)  # backtick quoted identifiers
                ]):
                    continue

                # Clean up the token
                p = re.sub(r"\s+as\s+.*$", '', p, flags=re.I)  # remove aliases
                p = p.split()[-1]  # get last part after spaces
                p = p.strip('`"\'')  # strip quotes/backticks
                
                # More lenient validation: check if it exists in any table's columns
                if p and p.lower() not in sql_keywords and not any(
                    p.lower() in {col.lower() for col in cols} 
                    for cols in table_columns.values()
                ):
                    unknown['unqualified_columns'].append(p)

        # If no unknowns, success
        if not unknown['table_column'] and not unknown['unqualified_columns']:
            return True, {}
        return False, unknown
    except Exception as e:
        logging.warning(f"Error validating SQL against schema: {e}")
        # On error be permissive but report problem
        return False, {'error': str(e)}

@main_routes.route('/config/db', methods=['GET', 'POST'])
def handle_db_config():
    if request.method == 'GET':
        return jsonify(Config.get_db_config())
    if request.method == 'POST':
        try:
            new_config = request.get_json(force=True)
        except Exception as e:
            return jsonify({'error': f'Invalid JSON: {str(e)}'}), 400

        required_fields = ['host', 'port', 'database', 'user', 'password']
        missing = [f for f in required_fields if not new_config.get(f)]
        if missing:
            return jsonify({'error': f'Missing required database configuration fields: {", ".join(missing)}'}), 400

        # Test the connection before saving
        if mysql_connector is None:
            return jsonify({'error': "Database connector not available on server; install 'mysql-connector-python' or 'PyMySQL'"}), 500
        try:
            conn = mysql_connector.connect(**new_config)
            conn.close()
        except Exception as e:
            return jsonify({'error': f'Database connection failed: {str(e)}'}), 400

        # Save DB config first
        try:
            Config.save_db_config(new_config)
        except Exception as e:
            return jsonify({'error': f'Failed to save DB config: {str(e)}'}), 500

        # === Clear ALL previous environment artifacts before bootstrapping new DB ===
        try:
            logging.info("Clearing all vector stores and saved data for database change...")
            # ...existing code for clearing vector stores, cache, etc...
            logging.info("Clearing all vector stores and saved data for database change...")
            
            # 1) Clear all vector stores and their cached data
            try:
                # Clear main vector store directory
                vs_root = Path(os.path.join(os.getcwd(), 'vector_store'))
                if vs_root.exists():
                    shutil.rmtree(str(vs_root))
                    os.makedirs(str(vs_root))
                    logging.info("Cleared main vector store directory")

                # Also clear vector stores in src directory if they exist
                src_vs_root = Path(os.path.join(os.path.dirname(__file__), '..', 'vector_store'))
                if src_vs_root.exists():
                    shutil.rmtree(str(src_vs_root))
                    os.makedirs(str(src_vs_root))
                    logging.info("Cleared src vector store directory")

            except Exception as e:
                logging.error(f"Error clearing vector stores: {e}")

            # Clear RAG manager cache
            try:
                rag_manager = RAGManager()
                # Force clear any cached data
                rag_manager.clear_all_data()
                # Reset schema context
                rag_manager.schema_context = SchemaContextManager()
                logging.info("Cleared RAG manager cache and schema context")
            except Exception as e:
                logging.error(f"Error clearing RAG manager cache: {e}")

            # 2) Clear external sources and schema knowledge
            try:
                # Clear all external sources
                sources_file = os.path.join(Config.CONFIG_DIR, 'sources.json')
                if os.path.exists(sources_file):
                    with open(sources_file, 'w') as f:
                        json.dump({'sources': []}, f)
                
                # Clear schema knowledge
                schema_file = os.path.join(Config.CONFIG_DIR, 'schema.json')
                if os.path.exists(schema_file):
                    with open(schema_file, 'w') as f:
                        json.dump({'tables': [], 'relationships': []}, f)

                logging.info("Cleared external sources and schema knowledge")
            except Exception as e:
                logging.error(f"Error clearing external sources and schema: {e}")

            # 3) Clear JSON config files
            try:
                # Clear guiders
                if os.path.exists(Config.GUIDERS_FILE):
                    with open(Config.GUIDERS_FILE, 'w') as f:
                        json.dump({}, f)
                        
                # Clear query history
                if os.path.exists(Config.HISTORY_FILE):
                    with open(Config.HISTORY_FILE, 'w') as f:
                        json.dump([], f)
                        
                # Clear context sources
                ctx_default = {'urls': [], 'documents': []}
                os.makedirs(os.path.dirname(Config.CONTEXT_FILE), exist_ok=True)
                with open(Config.CONTEXT_FILE, 'w') as f:
                    json.dump(ctx_default, f)
                    
                # Clear examples
                examples_file = os.path.join(Config.CONFIG_DIR, 'examples.json')
                if os.path.exists(examples_file):
                    with open(examples_file, 'w') as f:
                        json.dump({'examples': []}, f)
                        
                logging.info("Cleared all JSON configuration files")
            except Exception as e:
                logging.error(f"Error clearing JSON files: {e}")

            # 3) Remove uploaded files
            try:
                upload_dir = os.path.join(current_app.static_folder, 'uploads')
                if os.path.exists(upload_dir):
                    shutil.rmtree(upload_dir)
                    os.makedirs(upload_dir)
                    logging.info("Cleared uploads directory")
            except Exception as e:
                logging.error(f"Error clearing uploads: {e}")

            # 4) Clear business rules
            try:
                if os.path.exists(Config.GUIDERS_FILE):
                    with open(Config.GUIDERS_FILE, 'w') as f:
                        json.dump({}, f)
            except Exception as e:
                logging.warning(f"Failed to clear guiders file: {e}")

            # 2) Clear context sources (sources, uploaded docs list)
            try:
                ctx_default = {'urls': [], 'documents': []}
                os.makedirs(os.path.dirname(Config.CONTEXT_FILE), exist_ok=True)
                with open(Config.CONTEXT_FILE, 'w') as f:
                    json.dump(ctx_default, f)
            except Exception as e:
                logging.warning(f"Failed to clear context sources file: {e}")

            # 3) Remove vector store data for schema and business_rules
            try:
                vs_root = Path(os.path.join(os.getcwd(), 'vector_store'))
                for sub in ('schema', 'business_rules'):
                    p = vs_root / sub
                    if p.exists():
                        shutil.rmtree(str(p))
                        # recreate empty dir so RAG can reinitialize
                        p.mkdir(parents=True, exist_ok=True)
            except Exception as e:
                logging.warning(f"Failed to clear vector store directories: {e}")

            # 4) Optionally clear any persisted RAG index or cache files (top-level vector_store files)
            try:
                # remove any top-level artifacts in vector_store/<collection> if present
                # already handled above; keep lightweight
                pass
            except Exception:
                pass

            logging.info('Cleared previous guiders, context sources and vector store artifacts')
        except Exception:
            logging.exception('Error while clearing previous environment artifacts')

        # Try to bootstrap RAG schema knowledge from the newly saved DB config
        try:
            conn = get_db_connection()
            if conn:
                try:
                    rag = RAGManager()
                    rag.set_db_context(conn)
                    booted = rag.bootstrap_from_db(conn)
                    if booted:
                        return jsonify({'message': 'Database configuration updated and RAG bootstrapped successfully'})
                    else:
                        return jsonify({'message': 'Database configuration updated but RAG bootstrap returned no data'})
                finally:
                    try:
                        conn.close()
                    except Exception:
                        pass
        except Exception as e:
            logging.warning(f"Saved DB config but failed to bootstrap RAG: {e}")

        return jsonify({'message': 'Database configuration updated successfully'})


@main_routes.route('/config/smtp', methods=['GET', 'POST'])
def handle_smtp_config():
    """Get or save SMTP configuration used for sending emails."""
    if request.method == 'GET':
        try:
            cfg = Config.get_smtp_config()
            # Hide password when returning
            safe = cfg.copy() if isinstance(cfg, dict) else {}
            if safe.get('password'):
                safe['password'] = '(saved)'
            return jsonify(safe)
        except Exception as e:
            logging.error(f"Error getting SMTP config: {e}")
            return jsonify({'error': 'Failed to load SMTP config'}), 500

    # POST - save
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 415

    payload = request.get_json()
    # Basic allowed keys
    allowed = ['host', 'port', 'username', 'password', 'use_tls', 'from_address']
    cfg = {k: payload.get(k) for k in allowed if k in payload}

    # Normalize types
    if 'port' in cfg:
        try:
            cfg['port'] = int(cfg['port'])
        except Exception:
            pass
    if 'use_tls' in cfg:
        v = cfg['use_tls']
        if isinstance(v, str):
            cfg['use_tls'] = v.lower() in ('1', 'true', 'yes')

    ok = Config.save_smtp_config(cfg)
    if ok:
        # Also persist SMTP settings into a .env file so CLI fallbacks (scripts/send_email.py)
        # and other processes that read environment variables will pick them up.
        try:
            env_path = Path(os.path.join(os.getcwd(), '.env'))
            env = {}
            if env_path.exists():
                try:
                    with open(env_path, 'r', encoding='utf-8') as f:
                        for line in f:
                            line = line.strip()
                            if not line or line.startswith('#') or '=' not in line:
                                continue
                            k, v = line.split('=', 1)
                            env[k.strip()] = v.strip().strip('"').strip("'")
                except Exception:
                    env = {}

            # Map our config keys to SMTP_* env vars
            if 'host' in cfg and cfg.get('host'):
                env['SMTP_HOST'] = str(cfg.get('host'))
            if 'port' in cfg and cfg.get('port'):
                env['SMTP_PORT'] = str(cfg.get('port'))
            if 'username' in cfg and cfg.get('username'):
                env['SMTP_USER'] = str(cfg.get('username'))
            # Persist password to .env only if app setting allows it
            try:
                app_settings = Config.get_app_settings()
                persist_pass = bool(app_settings.get('persist_smtp_password'))
            except Exception:
                persist_pass = True
            if persist_pass and 'password' in cfg and cfg.get('password'):
                env['SMTP_PASS'] = str(cfg.get('password'))
            if 'use_tls' in cfg:
                env['SMTP_TLS'] = '1' if cfg.get('use_tls') else '0'
            if 'from_address' in cfg and cfg.get('from_address'):
                env['SMTP_FROM'] = str(cfg.get('from_address'))

            # Write back .env (overwrite existing keys)
            try:
                with open(env_path, 'w', encoding='utf-8') as f:
                    for k, v in env.items():
                        f.write(f"{k}={v}\n")
            except Exception as e:
                logging.warning(f"Failed to write .env with SMTP settings: {e}")

        except Exception as e:
            logging.warning(f"Error persisting SMTP to .env: {e}")

        return jsonify({'message': 'SMTP configuration saved successfully'})
    return jsonify({'error': 'Failed to save SMTP configuration'}), 500

@main_routes.route('/history', methods=['GET', 'POST', 'DELETE'])
def handle_history():
    if request.method == 'GET':
        try:
            history = HistoryManager.load_history()
            return jsonify(history)
        except Exception as e:
            logging.error(f"Error getting query history: {str(e)}")
            return jsonify({'error': 'Failed to load query history'}), 500
            
    elif request.method == 'POST':
        try:
            data = request.get_json()
            if not data or 'question' not in data:
                return jsonify({'error': 'Missing required fields'}), 400
                
            new_item = HistoryManager.add_item(
                question=data['question'],
                sql=data.get('sql', ''),
                result=data.get('result', ''),
                status=data.get('status', 'success')
            )
            
            if new_item:
                return jsonify(new_item)
            return jsonify({'error': 'Failed to save history item'}), 500
            
        except Exception as e:
            logging.error(f"Error saving to history: {str(e)}")
            return jsonify({'error': 'Failed to save to history'}), 500
            
    elif request.method == 'DELETE':
        try:
            if HistoryManager.clear_history():
                return jsonify({'message': 'Query history cleared successfully'})
            return jsonify({'error': 'Failed to clear history'}), 500
            
        except Exception as e:
            logging.error(f"Error clearing history: {str(e)}")
            return jsonify({'error': 'Failed to clear history'}), 500
            
@main_routes.route('/history/<string:id>', methods=['GET', 'DELETE'])
def handle_history_item(id):
    try:
        if request.method == 'GET':
            item = HistoryManager.get_item(id)
            if item:
                return jsonify(item)
            return jsonify({'error': 'History item not found'}), 404
            
        elif request.method == 'DELETE':
            if HistoryManager.delete_item(id):
                return jsonify({'message': 'History item deleted successfully'})
            return jsonify({'error': 'History item not found'}), 404
            
    except Exception as e:
        logging.error(f"Error handling history item: {str(e)}")
        return jsonify({'error': f'Failed to handle history item: {str(e)}'}), 500

@main_routes.route('/query', methods=['POST'])
def handle_query():
    """Handle queries and commands including immediate actions."""
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    command = data.get('question') or data.get('command')
    
    logging.info(f"Handling command/query: {command}")
    
    if not command:
        return jsonify({'error': 'Please provide a question or command'}), 400

    try:
        # Initialize command handler
        llm_client = LLMClient()
        query_handler = get_query_handler()
        command_handler = CommandHandler(llm_client, query_handler)

        # Include optional conversation history passed from client to provide context
        conversation = data.get('conversation') if isinstance(data, dict) else None
        # Allow the client to pass back clarifier selections via 'selected_tables'
        selected_tables = data.get('selected_tables') if isinstance(data, dict) else None
        result = command_handler.handle_command(command, conversation=conversation, selected_tables=selected_tables)

        # Determine status for history: previews are stored as 'preview'
        status = 'error' if result.get('error') else ('preview' if result.get('type') in ('query_preview', 'action_preview') else 'success')

        # Log to history (store preview entries so confirm can reference them)
        history_entry = {
            'id': str(uuid.uuid4()),
            'timestamp': datetime.now().isoformat(),
            'command': command,
            'type': result.get('type'),
            'sql': result.get('sql', ''),
            'explanation': result.get('explanation', ''),
            'status': status,
            'error': result.get('error', '')
        }

        history = Config.get_query_history()
        history.insert(0, history_entry)
        Config.save_query_history(history[:50])  # Keep last 50 queries

        # Return result plus history id so clients can confirm against it
        result['_history_id'] = history_entry['id']
        return jsonify(result)

    except Exception as e:
        error_message = str(e)
        logging.error(f"Error handling command: {error_message}")

        # Log failed command to history
        history_entry = {
            'id': str(uuid.uuid4()),
            'timestamp': datetime.now().isoformat(),
            'command': command,
            'status': 'error',
            'error': error_message
        }
        
        history = Config.get_query_history()
        history.insert(0, history_entry)
        Config.save_query_history(history[:50])
        
        return jsonify({'error': error_message}), 500


@main_routes.route('/confirm', methods=['POST'])
def confirm_action():
    """Execute a confirmed action or SQL (user clicked Confirm)."""
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    command = data.get('command')
    sql = data.get('sql')
    action = data.get('action')
    history_id = data.get('history_id')

    # We intentionally allow confirmation without a valid history_id so the
    # UI can execute whatever is currently shown on screen. Prefer to attach
    # to an existing preview entry if available, otherwise create a temporary
    # history entry so the execution is tracked.

    try:
        # Load history and attach to preview entry if possible
        history = Config.get_query_history()
        preview_entry = None
        if history_id:
            preview_entry = next((h for h in history if h.get('id') == history_id), None)

        # If no preview entry found, try to match by SQL or command text to reuse an entry
        if not preview_entry:
            if sql:
                preview_entry = next((h for h in history if h.get('sql') == sql), None)
            if not preview_entry and command:
                preview_entry = next((h for h in history if (h.get('command') or h.get('question') or '') == command), None)

        # If still not found, create a temporary history entry so execution is tracked
        if not preview_entry:
            preview_entry = {
                'id': str(uuid.uuid4()),
                'timestamp': datetime.now().isoformat(),
                'command': command or '',
                'sql': sql or '',
                'explanation': '',
                'status': 'preview',
                'error': ''
            }
            history.insert(0, preview_entry)
            Config.save_query_history(history[:50])

        # Proceed to execute, with server-side SQL validation to avoid executing queries
        query_handler = get_query_handler()
        command_handler = CommandHandler(None, query_handler)

        # Mark history as executing (always allowed)
        preview_entry['status'] = 'executing'
        Config.save_query_history(history[:50])

        # If action and sql provided (user edited SQL), prefer using provided sql
        if action and sql:
            # validate provided sql first
            valid = query_handler.validate_sql_against_schema(sql)
            if not valid.get('ok'):
                analysis = query_handler.analyze_error(sql, f"Validation failed: {valid}")
                preview_entry['status'] = 'error'
                preview_entry['error'] = f"SQL validation failed: {valid}"
                Config.save_query_history(history[:50])
                return jsonify({'error': 'SQL validation failed', 'details': valid, 'analysis': analysis}), 400
            res = command_handler._execute_action_with_optional_sql(command or '', action, sql)
        elif sql:
            # raw SQL provided — validate before executing
            valid = query_handler.validate_sql_against_schema(sql)
            if not valid.get('ok'):
                analysis = query_handler.analyze_error(sql, f"Validation failed: {valid}")
                preview_entry['status'] = 'error'
                preview_entry['error'] = f"SQL validation failed: {valid}"
                Config.save_query_history(history[:50])
                return jsonify({'error': 'SQL validation failed', 'details': valid, 'analysis': analysis}), 400
            res = {'type': 'sql_executed', 'results': query_handler.execute_sql(sql)}
        else:
            # Fallback: produce a preview first, validate generated SQL, then execute
            preview = command_handler.handle_command(command, execute=False, conversation=data.get('conversation'))
            generated_sql = preview.get('sql')
            if generated_sql:
                valid = query_handler.validate_sql_against_schema(generated_sql)
                if not valid.get('ok'):
                    analysis = query_handler.analyze_error(generated_sql, f"Validation failed: {valid}")
                    preview_entry['status'] = 'error'
                    preview_entry['error'] = f"SQL validation failed: {valid}"
                    Config.save_query_history(history[:50])
                    return jsonify({'error': 'Generated SQL validation failed', 'details': valid, 'analysis': analysis}), 400

            # Now execute for real
            res = command_handler.handle_command(command, execute=True, conversation=data.get('conversation'))

        # If SQL executed and returned tabular results, attempt to format using the same LLM
        try:
            formatted = None
            results = None
            # Normalize results extraction depending on command handler return shape
            if isinstance(res, dict) and res.get('results'):
                results = res.get('results')
            elif isinstance(res, dict) and res.get('type') == 'sql_executed' and res.get('results'):
                results = res.get('results')

            if results and isinstance(results, list):
                # Determine requested output format from the command or explanation text
                def detect_format(texts):
                    texts = texts or ''
                    # Look for format(...) pattern first
                    m = re.search(r"format\s*\(([^)]+)\)", texts, flags=re.I)
                    if m:
                        return m.group(1).strip().lower()
                    # Look for 'format: x' or 'as a X' patterns
                    m = re.search(r"format\s*[:\-]\s*(\w+)", texts, flags=re.I)
                    if m:
                        return m.group(1).strip().lower()
                    m = re.search(r"as an?\s+(tabular|chart|pdf|excel|image|csv|table)", texts, flags=re.I)
                    if m:
                        return m.group(1).strip().lower()
                    return 'tabular'

                requested_format = detect_format(preview_entry.get('command') or '')
                # Prepare a concise formatting prompt for the LLM
                sample_rows = results[:200]  # limit rows
                prompt = (
                    f"You are a formatting assistant. Convert the following SQL query results into the requested output format.\n"
                    f"Requested format: {requested_format}\n"
                    f"Return ONLY valid JSON with keys: format, content, content_type.\n"
                    f"- format: the chosen format name (tabular/chart/pdf/excel/image/csv).\n"
                    f"- content: for tabular return an HTML table string. For chart return a JSON object describing the chart (type, labels, datasets). For csv return CSV text. For pdf/excel/image, return a base64-encoded file content string and provide a filename in content_type.\n"
                    f"Here are the results (list of objects):\n{json.dumps(sample_rows, default=str, indent=2)}\n\n"
                    f"If you cannot produce the requested binary format (pdf/excel/image), instead return a CSV in 'content' and set format to 'csv' and content_type to 'text/csv'.\n"
                    f"Ensure the JSON is parseable."
                )

                try:
                    llm = LLMClient()
                    llm_resp = llm.generate(prompt)
                    # llm_resp should be text containing JSON
                    text = llm_resp if isinstance(llm_resp, str) else str(llm_resp)
                    # Extract JSON substring
                    s = text.find('{')
                    e = text.rfind('}')
                    if s != -1 and e != -1:
                        json_text = text[s:e+1]
                        try:
                            formatted = json.loads(json_text)
                        except Exception:
                            formatted = None
                    else:
                        formatted = None
                except Exception as e:
                    logging.warning(f"LLM formatting failed: {e}")

                # Fallback: render simple HTML table
                if not formatted:
                    # Build HTML table from results
                    if len(results) > 0:
                        cols = list(results[0].keys())
                    else:
                        cols = []
                    rows_html = []
                    rows_html.append('<table class="table table-sm table-striped"><thead><tr>')
                    for c in cols:
                        rows_html.append(f"<th>{html.escape(str(c))}</th>")
                    rows_html.append('</tr></thead><tbody>')
                    for r in results:
                        rows_html.append('<tr>')
                        for c in cols:
                            val = r.get(c, '')
                            rows_html.append(f"<td>{html.escape(str(val))}</td>")
                        rows_html.append('</tr>')
                    rows_html.append('</tbody></table>')
                    formatted = {
                        'format': 'tabular',
                        'content': ''.join(rows_html),
                        'content_type': 'text/html'
                    }

            # Attach formatted output to result for the client to render
            if formatted:
                res['formatted'] = formatted

        except Exception as e:
            logging.warning(f"Error formatting results with LLM: {e}")

        preview_entry['status'] = 'success' if not res.get('error') else 'error'
        preview_entry['result'] = res
        Config.save_query_history(history[:50])

        return jsonify(res)

    except Exception as e:
        logging.error(f"Error executing confirmed action: {e}")
        # Update history as error
        try:
            history = Config.get_query_history()
            preview_entry = next((h for h in history if h.get('id') == history_id), None)
            if preview_entry:
                preview_entry['status'] = 'error'
                preview_entry['error'] = str(e)
                Config.save_query_history(history[:50])
        except Exception:
            pass
        return jsonify({'error': str(e)}), 500

@main_routes.route('/execute', methods=['POST'])
def execute_sql():
    """Execute a raw SQL query sent from the client."""
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 415

    data = request.get_json()
    sql = data.get('sql')
    if not sql:
        return jsonify({'error': 'SQL is required'}), 400

    try:
        # Strong schema-aware validation: ensure SQL references real columns/tables
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'No DB connection available for validation'}), 500
        try:
            ok, details = validate_sql_against_rag_schema(sql, conn)
            if not ok:
                return jsonify({'error': 'SQL references unknown columns/tables', 'details': details}), 400
        finally:
            try: conn.close()
            except Exception: pass

        query_handler = get_query_handler()
        results = query_handler.execute_sql(sql)
        return jsonify(results)
    except Exception as e:
        logging.error(f"Error executing SQL: {str(e)}")
        return jsonify({'error': str(e)}), 500



# Unified /db/tables endpoint: always returns list of objects with table_name
@main_routes.route('/db/tables', methods=['GET'])
def list_db_tables():
    """Return a list of table names available in the configured database as objects with table_name."""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Could not connect to database'}), 500
        try:
            fetcher = SchemaFetcher(conn)
            tables = fetcher.get_tables() or []
            table_objs = []
            for t in tables:
                if isinstance(t, dict) and 'table_name' in t:
                    table_objs.append({'table_name': t['table_name']})
                else:
                    table_objs.append({'table_name': str(t)})
            return jsonify({'tables': table_objs})
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        logging.error(f"Error listing tables: {e}")
        return jsonify({'error': str(e)}), 500


@main_routes.route('/db/table/<string:table_name>/meta', methods=['GET'])
def table_meta(table_name):
    """Return columns, indexes, foreign keys and a brief description for a table."""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'No DB connection'}), 500
        try:
            fetcher = SchemaFetcher(conn)
            # Validate table exists
            tables = [t.get('table_name') for t in fetcher.get_tables() or []]
            if table_name not in tables:
                return jsonify({'error': 'Table not found'}), 404

            cols = fetcher.get_table_columns(table_name) or []

            # Indexes
            idx_cursor = conn.cursor(dictionary=True)
            try:
                idx_cursor.execute(f"SHOW INDEX FROM `{table_name}`")
                indexes = idx_cursor.fetchall() or []
            finally:
                idx_cursor.close()

            # Foreign keys (normalize to dicts)
            raw_fks = fetcher.get_foreign_keys(table_name) or []
            fks = []
            for row in raw_fks:
                # row might be tuple (COLUMN_NAME, REFERENCED_TABLE_NAME, REFERENCED_COLUMN_NAME)
                try:
                    fks.append({
                        'COLUMN_NAME': row[0],
                        'REFERENCED_TABLE_NAME': row[1],
                        'REFERENCED_COLUMN_NAME': row[2]
                    })
                except Exception:
                    # fallback if already dict-like
                    try:
                        fks.append(dict(row))
                    except Exception:
                        pass

            # Brief info via _get_table_info helper in SchemaFetcher if available
            brief = ''
            try:
                brief = fetcher._get_table_info(conn.cursor(dictionary=True), table_name) or ''
            except Exception:
                brief = ''

            return jsonify({
                'columns': cols,
                'indexes': indexes,
                'foreign_keys': fks,
                'brief': brief
            })
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        logging.error(f"Error fetching table meta for {table_name}: {e}")
        return jsonify({'error': str(e)}), 500


@main_routes.route('/db/table/<string:table_name>/rows', methods=['GET'])
def table_rows(table_name):
    """Return paginated rows from a table. Query params: page, page_size"""
    try:
        page = int(request.args.get('page', '1'))
        page_size = int(request.args.get('page_size', '25'))
        if page < 1: page = 1
        if page_size < 1: page_size = 25

        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'No DB connection'}), 500
        try:
            fetcher = SchemaFetcher(conn)
            tables = [t.get('table_name') for t in fetcher.get_tables() or []]
            if table_name not in tables:
                return jsonify({'error': 'Table not found'}), 404

            # Count total
            c = conn.cursor()
            try:
                c.execute(f"SELECT COUNT(*) FROM `{table_name}`")
                total = c.fetchone()[0]
            finally:
                c.close()

            offset = (page - 1) * page_size
            cur = conn.cursor(dictionary=True)
            try:
                cur.execute(f"SELECT * FROM `{table_name}` LIMIT %s OFFSET %s", (page_size, offset))
                rows = cur.fetchall() or []
            finally:
                cur.close()

            # Encode any binary/blob fields to base64 so JSON serialization won't fail
            try:
                for r in rows:
                    # r is a dict (dictionary=True) - convert byte-like values
                    for k, v in list(r.items()):
                        if v is None:
                            continue
                        try:
                            # memoryview support
                            if isinstance(v, memoryview):
                                b = v.tobytes()
                                r[k] = {"__b64__": True, "data": base64.b64encode(b).decode('ascii')}
                            elif isinstance(v, (bytes, bytearray)):
                                r[k] = {"__b64__": True, "data": base64.b64encode(bytes(v)).decode('ascii')}
                        except Exception:
                            # If conversion fails, stringify as fallback but still mark as base64-like
                            try:
                                r[k] = {"__b64__": True, "data": str(v)}
                            except Exception:
                                r[k] = None
            except Exception:
                # Non-fatal: if encoding fails, proceed with original rows (JSON may still fail if unhandled types present)
                pass

            return jsonify({'rows': rows, 'total': total})
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        logging.error(f"Error fetching rows for {table_name}: {e}")
        return jsonify({'error': str(e)}), 500


@main_routes.route('/db/table/<string:table_name>/lookup', methods=['GET'])
def table_lookup(table_name):
    """Lookup display values for a set of primary keys in another table.
    Query params:
      pk - primary key column name in referenced table (required)
      ids - comma-separated ids to lookup (required)
      display - optional display column to return (defaults to pk)
    Returns: { mappings: { id: display_value, ... } }
    """
    try:
        pk = request.args.get('pk')
        ids_param = request.args.get('ids')
        display = request.args.get('display')
        if not pk or not ids_param:
            return jsonify({'error': 'pk and ids are required'}), 400

        ids = [i for i in [s.strip() for s in ids_param.split(',')] if i != '']
        if not ids:
            return jsonify({'mappings': {}})

        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'No DB connection'}), 500
        try:
            # sanitize identifiers (very lightly) - ensure alphanum + underscore
            if not re.match(r'^[A-Za-z0-9_]+$', pk):
                return jsonify({'error': 'Invalid pk name'}), 400
            if not re.match(r'^[A-Za-z0-9_]+$', table_name):
                return jsonify({'error': 'Invalid table name'}), 400
            display_col = display if display and re.match(r'^[A-Za-z0-9_]+$', display) else pk

            # Build parameterized IN clause - use placeholders
            placeholders = ','.join(['%s'] * len(ids))
            cur = conn.cursor()
            try:
                sql = f"SELECT `{pk}`, `{display_col}` FROM `{table_name}` WHERE `{pk}` IN ({placeholders})"
                cur.execute(sql, tuple(ids))
                rows = cur.fetchall() or []
            finally:
                cur.close()

            mappings = {}
            for r in rows:
                try:
                    k = r[0]
                    v = r[1]
                    # ensure serializable
                    if isinstance(v, (bytes, bytearray, memoryview)):
                        try:
                            v = v.decode('utf-8', errors='ignore')
                        except Exception:
                            v = str(v)
                    mappings[str(k)] = v
                except Exception:
                    continue

            return jsonify({'mappings': mappings})
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        logging.error(f"Error in lookup for {table_name}: {e}")
        return jsonify({'error': str(e)}), 500

@main_routes.route('/query/preview', methods=['POST'])
def preview_query():
    data = request.get_json()
    question = data.get('question')
    if not question:
        return jsonify({'error': 'question is required'}), 400
    try:
        qh = get_query_handler()
        preview = qh.preview_query(question)
        return jsonify(preview)
    except Exception as e:
        logging.error(f"Error generating preview: {e}")
        return jsonify({'error': str(e)}), 500


@main_routes.route('/rag/bootstrap', methods=['POST'])
def rag_bootstrap():
    """Endpoint to manually trigger RAG bootstrap from the connected database."""
    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({
                'error': 'Could not connect to database for RAG bootstrap',
                'details': 'Check database connection settings and ensure the database is running.'
            }), 500
        try:
            rag = RAGManager()
            rag.set_db_context(conn)
            success = rag.bootstrap_from_db(conn)

            if success:
                knowledge = rag.get_all_knowledge()
                return jsonify({
                    'success': True,
                    'message': 'RAG bootstrap completed successfully',
                    'schema_count': len(knowledge.get('schema', [])),
                    'rules_count': len(knowledge.get('business_rules', []))
                })
            
            return jsonify({
                'success': False,
                'message': 'RAG bootstrap completed with warnings',
                'details': 'Some tables may have been skipped. Check server logs for details.'
            })
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        error_msg = str(e)
        logging.error(f"Error bootstrapping RAG: {error_msg}")
        return jsonify({
            'success': False,
            'error': 'Bootstrap process failed',
            'details': error_msg
        }), 500

@main_routes.route('/query/history', methods=['GET'])
def get_query_history():
    try:
        history = Config.get_query_history()
        return jsonify(history)
    except Exception as e:
        logging.error(f"Error fetching history: {e}")
        return jsonify({'error': 'Failed to load history'}), 500

@main_routes.route('/query/history/<string:query_id>', methods=['GET', 'PUT', 'DELETE'])
def manage_query_history(query_id):
    """Manage a single history item by its ID (string-safe)."""
    history = Config.get_query_history()
    if request.method == 'GET':
        item = next((h for h in history if str(h.get('id')) == str(query_id)), None)
        return jsonify(item) if item else ('', 404)

    elif request.method == 'PUT':
        updates = request.get_json()
        for i, h in enumerate(history):
            if str(h.get('id')) == str(query_id):
                history[i].update(updates)
                Config.save_query_history(history[:50])
                return ('', 200)
        return ('', 404)

    elif request.method == 'DELETE':
        # Remove any item whose id matches the provided query_id (string comparison)
        new_history = [h for h in history if not (str(h.get('id')) == str(query_id))]
        if len(new_history) == len(history):
            # nothing removed
            return ('', 404)
        Config.save_query_history(new_history[:50])
        return ('', 204)




@main_routes.route('/db/search', methods=['POST'])
def db_search():
    """Search text across multiple tables and return per-table matches.

    Request JSON: { query: 'pregnancy', tables: ['obs','person'] }
    If tables omitted or empty, searches all text-like columns in all tables.
    """
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 415
    payload = request.get_json()
    q = payload.get('query') or payload.get('q')
    tables = payload.get('tables') or []
    if not q:
        return jsonify({'error': 'query is required'}), 400

    try:
        conn = get_db_connection()
        if not conn:
            return jsonify({'error': 'Could not connect to database'}), 500

        try:
            fetcher = SchemaFetcher(conn)
            # determine which tables to search
            all_tables = [t.get('table_name') for t in fetcher.get_tables()]
            if not tables:
                search_tables = all_tables
            else:
                # sanitize provided table list by intersecting with actual tables
                search_tables = [t for t in tables if t in all_tables]

            results = []
            try:
                # For each table, identify text-like columns and search
                for table in search_tables:
                    text_cols = []
                    try:
                        cols = fetcher.get_table_columns(table)
                        for c in cols:
                            col_name = c.get('Field')
                            col_type = c.get('Type', '').lower()
                            if any(x in col_type for x in ('char', 'text', 'varchar')):
                                text_cols.append(col_name)
                    except Exception:
                        text_cols = []

                    table_result = {'table': table, 'matches': []}
                    if not text_cols:
                        # no searchable columns
                        continue

                    # Search each text column for occurrences and fetch full rows
                    for col in text_cols:
                        try:
                            dict_cursor = conn.cursor(dictionary=True)
                            pattern = f"%{q}%"
                            select_sql = f"SELECT * FROM `{table}` WHERE `{col}` LIKE %s LIMIT 5"
                            dict_cursor.execute(select_sql, (pattern,))
                            rows = dict_cursor.fetchall()
                            if rows:
                                samples = []
                                for r in rows:
                                    # convert bytes to strings if necessary
                                    simple = {}
                                    for k, v in r.items():
                                        if isinstance(v, (bytes, bytearray)):
                                            try:
                                                simple[k] = v.decode(errors='ignore')
                                            except Exception:
                                                simple[k] = str(v)
                                        else:
                                            simple[k] = v
                                    samples.append(simple)

                                # count total matches in that column
                                count_cursor = conn.cursor()
                                try:
                                    count_sql = f"SELECT COUNT(1) FROM `{table}` WHERE `{col}` LIKE %s"
                                    count_cursor.execute(count_sql, (pattern,))
                                    count = count_cursor.fetchone()[0]
                                finally:
                                    count_cursor.close()

                                table_result['matches'].append({'column': col, 'count': int(count), 'samples': samples})
                            dict_cursor.close()
                        except Exception as e:
                            logging.debug(f"Error searching {table}.{col}: {e}")

                    # Only include this table in results if we found matches
                    if table_result.get('matches'):
                        results.append(table_result)

            finally:
                pass

            return jsonify({'query': q, 'results': results})
        finally:
            try:
                conn.close()
            except Exception:
                pass

    except Exception as e:
        logging.error(f"Error performing DB search: {e}")
        return jsonify({'error': str(e)}), 500

@main_routes.route('/members/report', methods=['GET'])
def get_members_report():
    """Generate a report of new members with visualization."""
    try:
        # Get new members data
        members_query = """
        SELECT m.name AS new_member_name, 
               i.name AS invited_by_name 
        FROM users m 
        LEFT JOIN users i ON i.id = m.invited_by 
        WHERE m.status = 'New Members'
        """
        
        # Get status distribution for chart
        status_query = """
        SELECT status, COUNT(*) as member_count 
        FROM users 
        GROUP BY status
        """
        
        query_handler = get_query_handler()
        
        # Execute queries with enhanced validation
        try:
            members_data = query_handler.execute_sql(members_query)
            status_data = query_handler.execute_sql(status_query)
        except Exception as e:
            logging.error(f"Query execution error: {str(e)}")
            return jsonify({
                'error': 'Query execution failed',
                'details': str(e),
                'suggested_fix': 'Please check table and column names in your schema'
            }), 400

        # Prepare response with visualization hints for LLM
        response = {
            'data': {
                'table': {
                    'headers': ['New Member', 'Invited By'],
                    'rows': [
                        [row['new_member_name'], row['invited_by_name']] 
                        for row in members_data
                    ]
                },
                'chart': {
                    'type': 'column',
                    'title': 'Member Status Distribution',
                    'data': [
                        {
                            'category': row['status'],
                            'value': row['member_count']
                        } 
                        for row in status_data
                    ]
                }
            },
            'visualization': {
                'requirements': [
                    'Display a table showing new members and their inviters',
                    'Create a column chart showing the distribution of member status'
                ],
                'table_format': 'standard',
                'chart_type': 'column',
                'chart_options': {
                    'xAxis': 'Member Status',
                    'yAxis': 'Count',
                    'colors': ['#4e79a7', '#f28e2c']
                }
            }
        }
        
        return jsonify(response)
        
    except Exception as e:
        logging.error(f"Error generating members report: {str(e)}")
        return jsonify({'error': str(e)}), 500


@main_routes.route('/presentation/word', methods=['POST'])
def export_to_word():
    """Create a simple Word (.docx) document from provided JSON results and return it as attachment.

    Expected JSON: { title: str, summary: str (optional), table: { headers: [str], rows: [[val]] } }
    """
    if not request.is_json:
        return jsonify({'error': 'Request must be JSON'}), 415
    if Document is None:
        return jsonify({'error': 'python-docx is not installed on the server'}), 500

    data = request.get_json()
    title = data.get('title') or 'Export'
    summary = data.get('summary')
    table = data.get('table') or {}
    headers = table.get('headers') or []
    rows = table.get('rows') or []

    try:
        doc = Document()
        doc.add_heading(title, level=1)
        if summary:
            doc.add_paragraph(str(summary))

        if headers and rows:
            tbl = doc.add_table(rows=1, cols=len(headers))
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = str(h)
            for r in rows:
                row_cells = tbl.add_row().cells
                for i, cell_val in enumerate(r):
                    try:
                        row_cells[i].text = str(cell_val)
                    except Exception:
                        row_cells[i].text = ''

        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)

        filename = f"{title.replace(' ', '_')}.docx"
        return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        return jsonify({'error': f'Failed to create Word document: {e}'}), 500