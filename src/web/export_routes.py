from flask import Blueprint, request, send_file, jsonify
from ..config import Config
import pandas as pd
from io import BytesIO
import json
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import base64

export_routes = Blueprint('export', __name__, url_prefix='/api/export')

def create_excel(data):
    """Create Excel file from data"""
    output = BytesIO()
    if isinstance(data, dict) and 'results' in data:
        df = pd.DataFrame(data['results'])
    else:
        df = pd.DataFrame([data])
    
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='Results', index=False)
        
        # Add charts if present
        if 'chart' in data:
            chart_sheet = writer.book.add_worksheet('Charts')
            # Create charts based on data type
            if data['chart'].get('type') == 'bar':
                chart = writer.book.add_chart({'type': 'column'})
                # Add chart data and customize
                chart.add_series({
                    'values': f'=Results!$B$2:$B${len(df) + 1}',
                    'categories': f'=Results!$A$2:$A${len(df) + 1}',
                    'name': 'Values'
                })
                chart_sheet.insert_chart('A1', chart)
    
    output.seek(0)
    return output

def create_word(data):
    """Create Word document from data"""
    doc = Document()
    
    # Add title if present
    if 'title' in data:
        doc.add_heading(data['title'], 0)
    
    # Add explanation if present
    if 'explanation' in data:
        doc.add_paragraph(data['explanation'])
    
    # Add SQL if present
    if 'sql' in data:
        doc.add_heading('SQL Query', level=1)
        doc.add_paragraph(data['sql']).style = 'Code'
    
    # Add results table
    if 'results' in data:
        doc.add_heading('Results', level=1)
        table = doc.add_table(rows=1, cols=len(data['results'][0].keys()))
        table.style = 'Table Grid'
        
        # Add headers
        for i, key in enumerate(data['results'][0].keys()):
            table.cell(0, i).text = str(key)
        
        # Add data rows
        for row in data['results']:
            cells = table.add_row().cells
            for i, value in enumerate(row.values()):
                cells[i].text = str(value)
    
    # Add charts if present
    if 'chart' in data:
        doc.add_heading('Charts', level=1)
        # Create chart image
        plt.figure(figsize=(10, 6))
        if data['chart'].get('type') == 'bar':
            df = pd.DataFrame(data['results'])
            sns.barplot(data=df)
        plt.savefig('temp_chart.png')
        plt.close()
        
        doc.add_picture('temp_chart.png', width=Inches(6))
    
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def create_pdf(data):
    """Create PDF from data"""
    # Similar to Word but using a PDF library
    # This is a placeholder - implement with your preferred PDF library
    pass

@export_routes.route('/<format>', methods=['POST'])
def export_data(format):
    """Handle data export in various formats"""
    try:
        data = request.get_json()
        
        if format == 'excel':
            output = create_excel(data['data'])
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                as_attachment=True,
                download_name='results.xlsx'
            )
            
        elif format == 'word':
            output = create_word(data['data'])
            return send_file(
                output,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                as_attachment=True,
                download_name='results.docx'
            )
            
        elif format == 'pdf':
            # Implement PDF export
            return jsonify({'error': 'PDF export not yet implemented'}), 501
            
        elif format == 'csv':
            output = BytesIO()
            if isinstance(data['data'], dict) and 'results' in data['data']:
                df = pd.DataFrame(data['data']['results'])
            else:
                df = pd.DataFrame([data['data']])
            df.to_csv(output, index=False)
            output.seek(0)
            return send_file(
                output,
                mimetype='text/csv',
                as_attachment=True,
                download_name='results.csv'
            )
            
        else:
            return jsonify({'error': 'Unsupported format'}), 400
            
    except Exception as e:
        payload = Config.client_error_payload(str(e), e)
        return jsonify(payload), 500